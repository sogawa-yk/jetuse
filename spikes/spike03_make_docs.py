"""SPIKE-03: 日本語ダミー文書（社内規程風）3種を生成する。

  - 出張旅費規程 (PDF, reportlab + CIDフォント HeiseiKakuGo-W5)
  - 在宅勤務規程 (docx, python-docx)
  - 経費精算ガイドライン (md)

出力先: spikes/data/
"""
from pathlib import Path

OUT = Path(__file__).resolve().parent / "data"
OUT.mkdir(exist_ok=True)

TRAVEL = [
    ("第1条（目的）", "本規程は、株式会社サンプル商事（以下「会社」）の役職員が業務のため出張する場合の手続きおよび旅費の支給基準を定める。"),
    ("第2条（出張の定義）", "出張とは、従業員が所属事業所を起点として片道50キロメートル以上の地への業務遂行のための旅行をいう。片道50キロメートル未満は近距離出張とし、交通費実費のみを支給する。"),
    ("第3条（出張の承認）", "出張をしようとする者は、出発日の3営業日前までに出張申請システムにより所属長の承認を得なければならない。緊急の場合は口頭承認の後、事後速やかに申請を行う。"),
    ("第4条（旅費の種類）", "旅費は、交通費、日当、宿泊費および雑費とする。"),
    ("第5条（交通費）", "交通費は実費を支給する。新幹線は普通車指定席、航空機はエコノミークラスを原則とする。部長職以上は新幹線グリーン車の利用を認める。"),
    ("第6条（日当）", "日当は、国内出張の場合、一般職2,000円、管理職3,000円、役員5,000円を1日あたり支給する。海外出張の日当は別表に定める。"),
    ("第7条（宿泊費）", "宿泊費の上限は、一般職10,000円、管理職13,000円、役員18,000円とする。東京23区内および大阪市内での宿泊は、それぞれ2,000円を上限額に加算する。"),
    ("第8条（精算）", "出張者は、帰着後5営業日以内に経費精算システムにより精算を行い、領収書を添付しなければならない。"),
    ("第9条（改廃）", "本規程の改廃は、取締役会の決議による。附則: 本規程は2024年4月1日から施行する。"),
]

REMOTE = [
    ("第1条（目的）", "本規程は、株式会社サンプル商事における在宅勤務制度の適用条件および運用ルールを定めることを目的とする。"),
    ("第2条（適用対象）", "在宅勤務制度は、入社後6ヶ月を経過した正社員および契約社員に適用する。試用期間中の者、および業務の性質上出社が必須と所属長が判断した者は対象外とする。"),
    ("第3条（在宅勤務の上限）", "在宅勤務は週3日を上限とする。ただし、育児・介護の事由がある場合は所属長の承認により週5日まで認める。"),
    ("第4条（勤務時間）", "在宅勤務時の勤務時間は就業規則に定める所定労働時間とし、始業・終業時刻をチャットツールで所属長に報告する。中抜けは1日2時間まで認め、勤怠システムに記録する。"),
    ("第5条（通信費等の補助）", "在宅勤務手当として月額3,000円を支給する。在宅勤務日数が月4日未満の場合は支給しない。"),
    ("第6条（セキュリティ）", "在宅勤務時は会社貸与PCのみを使用し、公衆Wi-Fiの利用を禁止する。VPN接続を必須とし、家族を含む第三者に画面を閲覧させてはならない。"),
    ("第7条（光熱費）", "自宅の光熱費は在宅勤務手当に含まれるものとし、個別の精算は行わない。"),
]

EXPENSE_MD = """# 経費精算ガイドライン

株式会社サンプル商事 経理部（2025年10月改訂）

## 1. 基本ルール

- 経費精算は**発生月の翌月10日**までに経費精算システムで申請すること。
- 1件あたり**5,000円以上**の支出は領収書（宛名: 会社名義）が必須。5,000円未満はレシート可。
- 電子帳簿保存法対応のため、領収書はスキャンしてシステムにアップロードする。原本の提出は不要。

## 2. 接待交際費

- 1人あたり**10,000円以内**の飲食は会議費として処理できる。参加者全員の氏名・所属の記録が必要。
- 10,000円を超える接待は事前に部長承認を要する。
- 社内飲み会は福利厚生費の対象外（自己負担）。

## 3. 交通費

- 通勤定期区間内の移動は精算不可。
- タクシーは、終電後または重量物運搬時のみ利用可。利用理由をシステムに記載すること。

## 4. 立替払いの禁止事項

- 10万円を超える物品購入の立替は禁止。購買部経由で発注すること。
- 商品券・ギフトカード類の購入は経理部の事前承認が必要。

## 5. 問い合わせ

経理部 精算チーム（内線 1234 / keiri@example.co.jp）
"""


def make_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    title = ParagraphStyle("t", fontName="HeiseiKakuGo-W5", fontSize=16, spaceAfter=12)
    head = ParagraphStyle("h", fontName="HeiseiKakuGo-W5", fontSize=12, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("b", fontName="HeiseiKakuGo-W5", fontSize=10.5, leading=16)

    doc = SimpleDocTemplate(str(OUT / "travel-policy.pdf"), pagesize=A4)
    flow = [Paragraph("出張旅費規程（株式会社サンプル商事）", title), Spacer(1, 8)]
    for h, b in TRAVEL:
        flow += [Paragraph(h, head), Paragraph(b, body)]
    doc.build(flow)


def make_docx():
    import docx
    d = docx.Document()
    d.add_heading("在宅勤務規程（株式会社サンプル商事）", level=1)
    for h, b in REMOTE:
        d.add_heading(h, level=2)
        d.add_paragraph(b)
    d.save(OUT / "remote-work-policy.docx")


def make_md():
    (OUT / "expense-guidelines.md").write_text(EXPENSE_MD)


def make_txt():
    # docxがVector Store未対応の場合の代替（同内容のプレーンテキスト）
    body = "在宅勤務規程（株式会社サンプル商事）\n\n"
    body += "\n\n".join(f"{h}\n{b}" for h, b in REMOTE)
    (OUT / "remote-work-policy.txt").write_text(body)


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_md()
    make_txt()
    for p in sorted(OUT.iterdir()):
        print(p.name, p.stat().st_size, "bytes")
